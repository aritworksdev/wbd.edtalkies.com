from __future__ import annotations

from pathlib import Path

from docx import Document
import fitz

from aiboard_app.documents.document_text_extractor import DocumentTextExtractor
from aiboard_app.recognition.ocr_result import OcrResult


class FakeOcrService:
    def recognize(self, image_bytes: bytes) -> OcrResult:
        assert image_bytes
        return OcrResult("recognized image text", 0.88, "fake-ocr")

    def recognize_document(self, image_bytes: bytes) -> OcrResult:
        return self.recognize(image_bytes)


def test_extract_plain_text_document(tmp_path: Path) -> None:
    path = tmp_path / "lesson.txt"
    path.write_text("Photosynthesis converts light into chemical energy.", encoding="utf-8")

    extraction = DocumentTextExtractor(FakeOcrService()).extract(path)
    result = extraction.result

    assert "Photosynthesis" in result.text
    assert result.confidence == 1.0
    assert result.provider == "document-txt"


def test_extract_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "lesson.docx"
    document = Document()
    document.add_paragraph("Lesson heading")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Meaning"
    document.save(path)

    result = DocumentTextExtractor(FakeOcrService()).extract(path).result

    assert "Lesson heading" in result.text
    assert "Term\tMeaning" in result.text


def test_extract_html_without_script_content(tmp_path: Path) -> None:
    path = tmp_path / "lesson.html"
    path.write_text(
        "<h1>Gravity</h1><script>ignore()</script><p>Gravity attracts masses.</p>",
        encoding="utf-8",
    )

    result = DocumentTextExtractor(FakeOcrService()).extract(path).result

    assert result.text == "Gravity Gravity attracts masses."


def test_image_document_uses_ocr_service(tmp_path: Path) -> None:
    path = tmp_path / "board.png"
    path.write_bytes(b"image-data")

    extraction = DocumentTextExtractor(FakeOcrService()).extract(path)
    result = extraction.result

    assert result.text == "recognized image text"
    assert result.provider == "fake-ocr"
    assert extraction.google_vision_images == (b"image-data",)


def test_extract_pdf_embedded_text(tmp_path: Path) -> None:
    path = tmp_path / "lesson.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Newton described gravity.")
    document.save(path)
    document.close()

    extraction = DocumentTextExtractor(FakeOcrService()).extract(path)
    result = extraction.result

    assert "Newton described gravity." in result.text
    assert result.provider == "document-pdf"
    assert extraction.google_vision_images == ()


def test_mixed_pdf_preserves_text_and_ocrs_scanned_pages(tmp_path: Path) -> None:
    path = tmp_path / "mixed.pdf"
    document = fitz.open()
    text_page = document.new_page()
    text_page.insert_text((72, 72), "Digital first page.")
    document.new_page()
    document.save(path)
    document.close()

    extraction = DocumentTextExtractor(FakeOcrService()).extract(path)

    assert "Digital first page." in extraction.result.text
    assert "recognized image text" in extraction.result.text
    assert len(extraction.google_vision_images) == 1
