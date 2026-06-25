from __future__ import annotations

import html
import re
from dataclasses import dataclass
from pathlib import Path

from aiboard_app.recognition.ocr_result import OcrResult, OcrWord
from aiboard_app.recognition.ocr_service import OcrService


@dataclass(frozen=True)
class DocumentExtraction:
    result: OcrResult
    google_vision_images: tuple[bytes, ...] = ()


class DocumentTextExtractor:
    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
        ".md",
        ".markdown",
        ".html",
        ".htm",
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".tif",
        ".tiff",
    }

    def __init__(self, ocr_service: OcrService) -> None:
        self._ocr_service = ocr_service

    def extract(self, path: Path) -> DocumentExtraction:
        if not path.is_file():
            raise FileNotFoundError(path)
        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported document type: {extension or 'unknown'}")

        if extension == ".pdf":
            return self._extract_pdf(path)
        if extension == ".docx":
            return DocumentExtraction(self._extract_docx(path))
        if extension in {".html", ".htm"}:
            return DocumentExtraction(self._extract_html(path))
        if extension in {".txt", ".md", ".markdown"}:
            return DocumentExtraction(
                self._direct_result(
                    path.read_text(encoding="utf-8-sig"),
                    f"document-{extension[1:]}",
                )
            )
        image_bytes = path.read_bytes()
        return DocumentExtraction(
            self._ocr_service.recognize_document(image_bytes),
            (image_bytes,),
        )

    def _extract_pdf(self, path: Path) -> DocumentExtraction:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF extraction requires pypdf.") from exc

        reader = PdfReader(str(path))
        embedded_pages = [(page.extract_text() or "").strip() for page in reader.pages]

        # Mixed PDFs can contain both digital-text and scanned pages. Preserve
        # embedded text and OCR only the pages that do not contain usable text.
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("Scanned PDF OCR requires PyMuPDF.") from exc

        results: list[OcrResult] = []
        page_images: list[bytes] = []
        with fitz.open(path) as document:
            for index, page in enumerate(document):
                embedded = embedded_pages[index] if index < len(embedded_pages) else ""
                if embedded:
                    results.append(self._direct_result(embedded, "document-pdf"))
                    continue
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
                image_bytes = pixmap.tobytes("png")
                page_images.append(image_bytes)
                results.append(self._ocr_service.recognize_document(image_bytes))
        provider = "document-pdf" if not page_images else "document-pdf-mixed-ocr"
        return DocumentExtraction(self._combine_ocr_results(results, provider), tuple(page_images))

    @staticmethod
    def _extract_docx(path: Path) -> OcrResult:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX extraction requires python-docx.") from exc
        document = Document(str(path))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append("\t".join(cells))
        return DocumentTextExtractor._direct_result("\n".join(blocks), "document-docx")

    @staticmethod
    def _extract_html(path: Path) -> OcrResult:
        content = path.read_text(encoding="utf-8-sig")
        content = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", content)
        content = re.sub(r"(?s)<[^>]+>", " ", content)
        text = re.sub(r"\s+", " ", html.unescape(content)).strip()
        return DocumentTextExtractor._direct_result(text, "document-html")

    @staticmethod
    def _direct_result(text: str, provider: str) -> OcrResult:
        text = text.strip()
        words = tuple(OcrWord(token, 1.0) for token in text.split())
        return OcrResult(
            text=text,
            confidence=1.0 if text else 0.0,
            provider=provider,
            words=words,
            model_name="Embedded document text",
        )

    @staticmethod
    def _combine_ocr_results(results: list[OcrResult], provider: str) -> OcrResult:
        populated = [result for result in results if result.text.strip()]
        if not populated:
            return OcrResult.empty(provider)
        confidences = [result.confidence for result in populated if result.confidence is not None]
        confidence = sum(confidences) / len(confidences) if confidences else None
        return OcrResult(
            text="\n\n".join(result.text for result in populated),
            confidence=confidence,
            provider=provider,
            words=tuple(word for result in populated for word in result.words),
            attempts=tuple(attempt for result in populated for attempt in result.attempts),
            errors=tuple(error for result in populated for error in result.errors),
            model_name=", ".join(
                dict.fromkeys(result.model_name or result.provider for result in populated)
            ),
        )
