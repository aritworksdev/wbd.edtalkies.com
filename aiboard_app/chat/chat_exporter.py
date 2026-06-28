from __future__ import annotations

import re
import zipfile
from html import unescape
from xml.sax.saxutils import escape
from pathlib import Path

from aiboard_app.chat.chat_record import ChatRecord


class ChatExporter:
    def __init__(self, export_dir: Path) -> None:
        self._export_dir = export_dir / "chat_exports"

    def default_path(self, chat: ChatRecord, file_type: str) -> Path:
        stamp = chat.created_at.strftime("%Y%m%d_%H%M%S")
        suffix = file_type.lower().lstrip(".")
        return self._export_dir / f"AiBoard_Chat_{stamp}.{suffix}"

    def export(self, chat: ChatRecord, file_type: str, destination: Path | None = None) -> Path:
        suffix = file_type.lower().lstrip(".")
        path = destination or self.default_path(chat, suffix)
        path.parent.mkdir(parents=True, exist_ok=True)
        if suffix == "txt":
            self._export_txt(chat, path)
        elif suffix == "docx":
            self._export_docx(chat, path)
        elif suffix == "pdf":
            self._export_pdf(chat, path)
        elif suffix == "xlsx":
            self._export_xlsx(chat, path)
        elif suffix == "pptx":
            self._export_pptx(chat, path)
        else:
            raise ValueError(f"Unsupported export type: {file_type}")
        return path

    def _export_txt(self, chat: ChatRecord, path: Path) -> None:
        path.write_text(self._plain_content(chat), encoding="utf-8")

    def _export_docx(self, chat: ChatRecord, path: Path) -> None:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency is declared.
            raise RuntimeError("DOCX export requires python-docx.") from exc

        document = Document()
        document.add_heading("AiBoard Chat", level=1)
        document.add_paragraph(f"Created: {chat.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        document.add_paragraph(f"Chat ID: {chat.id}")
        if chat.session_id:
            document.add_paragraph(f"Session ID: {chat.session_id}")
        document.add_heading("Recognized Text / Request", level=2)
        document.add_paragraph(chat.request_text)
        document.add_heading(chat.response.title or "EdTalkies Response", level=2)
        document.add_paragraph(self._response_text(chat))
        document.save(path)

    def _export_pdf(self, chat: ChatRecord, path: Path) -> None:
        try:
            import fitz
        except ImportError as exc:  # pragma: no cover - dependency is declared.
            raise RuntimeError("PDF export requires PyMuPDF.") from exc

        doc = fitz.open()
        page = doc.new_page()
        margin = 54
        cursor_y = margin
        lines = self._plain_content(chat).splitlines()
        for line in lines:
            if cursor_y > page.rect.height - margin:
                page = doc.new_page()
                cursor_y = margin
            page.insert_text(
                (margin, cursor_y),
                line or " ",
                fontsize=11,
                fontname="helv",
                color=(0, 0, 0),
            )
            cursor_y += 16
        doc.save(path)
        doc.close()

    def _export_xlsx(self, chat: ChatRecord, path: Path) -> None:
        rows = [
            ("Field", "Value"),
            ("Created", chat.created_at.strftime("%Y-%m-%d %H:%M:%S")),
            ("Chat ID", chat.id),
            ("Recognized Text / Request", chat.request_text),
            ("Response Title", chat.response.title or "EdTalkies Response"),
            ("AI Response", self._response_text(chat)),
        ]
        sheet_rows = []
        for row_index, (field, value) in enumerate(rows, start=1):
            sheet_rows.append(
                "<row r=\"{row}\">"
                "<c r=\"A{row}\" t=\"inlineStr\"><is><t>{field}</t></is></c>"
                "<c r=\"B{row}\" t=\"inlineStr\"><is><t>{value}</t></is></c>"
                "</row>".format(
                    row=row_index,
                    field=escape(field),
                    value=escape(value),
                )
            )
        files = {
            "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>""",
            "_rels/.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>""",
            "xl/workbook.xml": """<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="AiBoard Chat" sheetId="1" r:id="rId1"/></sheets>
</workbook>""",
            "xl/_rels/workbook.xml.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>""",
            "xl/worksheets/sheet1.xml": f"""<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>{''.join(sheet_rows)}</sheetData>
</worksheet>""",
        }
        self._write_zip(path, files)

    def _export_pptx(self, chat: ChatRecord, path: Path) -> None:
        body = escape(self._response_text(chat))
        request = escape(chat.request_text)
        created = escape(chat.created_at.strftime("%Y-%m-%d %H:%M:%S"))
        files = {
            "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>
  <Override PartName="/ppt/slides/slide1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>
</Types>""",
            "_rels/.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
</Relationships>""",
            "ppt/presentation.xml": """<?xml version="1.0" encoding="UTF-8"?>
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst><p:sldId id="256" r:id="rId1"/></p:sldIdLst>
  <p:sldSz cx="12192000" cy="6858000"/>
</p:presentation>""",
            "ppt/_rels/presentation.xml.rels": """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide1.xml"/>
</Relationships>""",
            "ppt/slides/slide1.xml": f"""<?xml version="1.0" encoding="UTF-8"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld><p:spTree>
    <p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/>
    <p:sp><p:nvSpPr><p:cNvPr id="2" name="Title"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:t>AiBoard Chat</a:t></a:r></a:p></p:txBody></p:sp>
    <p:sp><p:nvSpPr><p:cNvPr id="3" name="Content"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:t>Created: {created}</a:t></a:r></a:p><a:p><a:r><a:t>Request: {request}</a:t></a:r></a:p><a:p><a:r><a:t>Response: {body}</a:t></a:r></a:p></p:txBody></p:sp>
  </p:spTree></p:cSld>
</p:sld>""",
        }
        self._write_zip(path, files)

    def _plain_content(self, chat: ChatRecord) -> str:
        return "\n".join(
            [
                "AiBoard Chat",
                f"Created: {chat.created_at.strftime('%Y-%m-%d %H:%M:%S')}",
                f"Chat ID: {chat.id}",
                "",
                "Recognized Text / Request:",
                chat.request_text,
                "",
                f"{chat.response.title or 'EdTalkies Response'}:",
                self._response_text(chat),
                "",
            ]
        )

    @staticmethod
    def _response_text(chat: ChatRecord) -> str:
        if chat.response.text.strip():
            return chat.response.text.strip()
        html = chat.response.html
        text = re.sub(r"(?i)<br\s*/?>", "\n", html)
        text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", text)
        text = re.sub(r"<[^>]+>", "", text)
        return unescape(text).strip()

    @staticmethod
    def _write_zip(path: Path, files: dict[str, str]) -> None:
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
            for name, content in files.items():
                archive.writestr(name, content)
